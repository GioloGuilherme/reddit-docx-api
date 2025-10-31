"""
FastAPI backend for Reddit thread to DOCX conversion.

This service exposes two endpoints:

* POST `/process` — Accepts Reddit API credentials and a text file with one
  Reddit thread URL per line. It fetches each thread, converts it into a
  Word document, zips all generated documents, and returns JSON containing
  per‑URL results along with a URL to download the ZIP file.

* GET `/download/{job_id}` — Serves the ZIP file generated during a
  `/process` request.

The core scraping and document generation logic mirrors functionality from
an existing Colab script. Comments are recursively flattened with
indentation preserved, media thumbnails are embedded when possible, and
upvotes are recorded. A 30‑second delay between Reddit API calls is
enforced to respect rate limits.

To deploy this service on Railway:
  1. Include this file and a requirements.txt in your project directory.
  2. Add a `railway.json` specifying the start command (e.g.
     `uvicorn app:app --host 0.0.0.0 --port $PORT`).
  3. Use the Railway CLI (`railway init` and `railway up`) to deploy.

This file intentionally contains no user credentials or Colab‑specific
dependencies. It uses FastAPI for the API, PRAW for Reddit access, and
python‑docx to generate documents.
"""

import os
import threading
import time
import uuid
import shutil
import zipfile
from datetime import datetime
from typing import List, Optional

import praw
import requests
from fastapi import FastAPI, UploadFile, Form, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import Pt, Inches
from urllib.parse import urlparse, parse_qs

# Base directory for storing job artefacts. Each API request creates a
# subfolder within this directory, containing generated DOCX files and the
# final ZIP archive. Railway's filesystem is persisted across deployments,
# so you may wish to periodically prune this directory in production.
BASE_JOBS_DIR = os.environ.get("JOBS_DIR", "./jobs")
DELAY_SECONDS = int(os.environ.get("DELAY_SECONDS", "30"))  # throttle between Reddit calls

# In‑memory store for job metadata. Each job_id key maps to a dict
# tracking the current status, total number of URLs, a list of per‑URL
# results, and the relative ZIP download path once available. In a
# production environment, you would likely persist this to a database
# or use a proper task queue.
jobs: dict[str, dict] = {}

os.makedirs(BASE_JOBS_DIR, exist_ok=True)

app = FastAPI()

# Allow cross‑origin requests so that a separate frontend (e.g. Lovable
# or another domain) can call this API. In production, replace "*"
# with your actual frontend origin.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def authenticate_reddit(client_id: str, client_secret: str, user_agent: str) -> praw.Reddit:
    """Authenticate against Reddit using provided credentials."""
    return praw.Reddit(
        client_id=client_id,
        client_secret=client_secret,
        user_agent=user_agent,
    )


def fetch_comments(comments, level: int = 0) -> List[tuple[str, str, int]]:
    """
    Recursively collect all comments into a flat list of tuples. Each
    tuple contains the indented comment body, author name, and upvote
    score. Replies are indented by four spaces per level.
    """
    comments_data: List[tuple[str, str, int]] = []
    for comment in comments:
        comment_text = comment.body
        comment_author = str(comment.author) if comment.author else "deleted"
        comment_upvotes = comment.score
        indentation = '    ' * level
        comments_data.append((indentation + comment_text, comment_author, comment_upvotes))
        if comment.replies:
            replies = fetch_comments(comment.replies, level + 1)
            comments_data.extend(replies)
    return comments_data


def get_image_or_youtube_thumbnail(url: str) -> Optional[str]:
    """
    Determine if the provided URL points to an image or YouTube video. If
    it's an image, return the URL. For YouTube links, construct a link to
    the video's thumbnail. Otherwise return None.
    """
    if url.endswith((
        '.jpg', '.jpeg', '.png', '.gif', '.webp'
    )):
        return url
    parsed_url = urlparse(url)
    if "youtube.com" in parsed_url.netloc or "youtu.be" in parsed_url.netloc:
        video_id: Optional[str] = None
        if "youtube.com" in parsed_url.netloc:
            video_id = parse_qs(parsed_url.query).get("v", [None])[0]
        else:
            # for youtu.be links, the path starts with '/<id>'
            video_id = parsed_url.path.lstrip('/')
        if video_id:
            return f"https://img.youtube.com/vi/{video_id}/0.jpg"
    return None


def get_reddit_thread_data(reddit: praw.Reddit, url: str) -> tuple[
    str, str, int, List[tuple[str, str, int]], List[str], str, str
]:
    """
    Given a Reddit thread URL, fetch relevant data: title, body, upvotes,
    flattened comments, list of media URLs (image or YouTube thumbnail),
    post date (as DD.MM.YYYY), and submission ID.
    """
    submission = reddit.submission(url=url)
    thread_title = submission.title
    thread_body = submission.selftext
    thread_upvotes = submission.score
    post_date = datetime.utcfromtimestamp(submission.created_utc).strftime("%d.%m.%Y")
    submission.comments.replace_more(limit=None)
    comments_data = fetch_comments(submission.comments)
    media_urls: List[str] = []
    if submission.url:
        media_url = get_image_or_youtube_thumbnail(submission.url)
        if media_url:
            media_urls.append(media_url)
    return (
        thread_title,
        thread_body,
        thread_upvotes,
        comments_data,
        media_urls,
        post_date,
        submission.id,
    )


def download_image(url: str, save_folder: str) -> Optional[str]:
    """
    Download an image from `url` into `save_folder`. Returns the file path
    or None on failure.
    """
    # Use only the basename of the URL path to avoid directory traversal.
    basename = os.path.basename(urlparse(url).path) or "image.jpg"
    filename = os.path.join(save_folder, basename)
    try:
        response = requests.get(url, stream=True, timeout=10)
        if response.status_code == 200:
            with open(filename, 'wb') as file:
                for chunk in response.iter_content(chunk_size=8192):
                    file.write(chunk)
            return filename
    except Exception:
        pass
    return None


def generate_docx(
    thread_url: str,
    thread_title: str,
    thread_body: str,
    thread_upvotes: int,
    comments_data: List[tuple[str, str, int]],
    media_urls: List[str],
    post_date: str,
    destination_folder: str,
    submission_id: str,
) -> str:
    """
    Generate a DOCX file representing a Reddit thread. The file includes the
    thread URL, title, body, upvote count, embedded media, and all
    comments. It returns the filename of the generated document.
    """
    doc = Document()
    # URL at the top in a small font
    url_paragraph = doc.add_paragraph()
    url_run = url_paragraph.add_run(thread_url)
    url_run.font.size = Pt(9)
    # Title as heading
    title_paragraph = doc.add_heading(thread_title, level=0)
    if title_paragraph.runs:
        title_paragraph.runs[0].bold = True
    # Body
    if thread_body:
        doc.add_paragraph(thread_body)
    # Upvotes
    doc.add_paragraph(f"Upvotes: {thread_upvotes}")
    # Media
    for media_url in media_urls:
        image_path = download_image(media_url, destination_folder)
        if image_path:
            doc.add_paragraph("Media from thread:")
            doc.add_picture(image_path, width=Inches(5))
    # Comments
    for comment, author, upvotes in comments_data:
        comment_para = doc.add_paragraph()
        comment_run = comment_para.add_run(comment)
        comment_run.italic = True
        author_para = doc.add_paragraph(f"- {author}")
        author_para.add_run(f" (Upvotes: {upvotes})")
        # Slight indent on author line
        author_para.paragraph_format.left_indent = Inches(0.1)
    # Save document
    os.makedirs(destination_folder, exist_ok=True)
    filename = f"{post_date}_{submission_id}.docx"
    output_path = os.path.join(destination_folder, filename)
    doc.save(output_path)
    return filename


def zip_folder(folder_path: str, zip_path: str) -> None:
    """Compress all DOCX files in `folder_path` into a ZIP archive at `zip_path`."""
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, _, files in os.walk(folder_path):
            for f in files:
                if f.lower().endswith(".docx"):
                    full = os.path.join(root, f)
                    zf.write(full, arcname=os.path.basename(full))


@app.post("/api/process")
async def process_threads(
    client_id: str = Form(...),
    client_secret: str = Form(...),
    user_agent: str = Form(...),
    urls_file: UploadFile = Form(...),
):
    """
    Accept Reddit credentials and a file containing thread URLs (one per line).
    Create a job folder, fetch each thread sequentially with a delay, generate
    corresponding DOCX files, and return a JSON response summarizing the
    results along with a download URL for the ZIP.
    """
    # Read URLs from uploaded file
    file_bytes = await urls_file.read()
    try:
        raw_text = file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        raise HTTPException(status_code=400, detail="Unable to decode uploaded file.")
    url_list = [u.strip() for u in raw_text.splitlines() if u.strip()]
    total = len(url_list)
    if total == 0:
        raise HTTPException(status_code=400, detail="No URLs found in uploaded file.")
    # Create a job ID and job directory
    job_id = str(uuid.uuid4())
    job_dir = os.path.join(BASE_JOBS_DIR, job_id)
    os.makedirs(job_dir, exist_ok=True)
    # Initialize job state
    jobs[job_id] = {
        "status": "pending",
        "total": total,
        "completed": 0,
        "results": [
            {"url": url, "status": "pending", "filename": None, "error": None}
            for url in url_list
        ],
        "zip_url": None,
    }

    def process_job():
        """
        Worker function to process a list of URLs into DOCX files and a ZIP
        archive. Updates the global `jobs` dict as it progresses.
        """
        try:
            # Authenticate to Reddit once at the start
            try:
                reddit = authenticate_reddit(client_id, client_secret, user_agent)
                # Validate credentials with a lightweight request
                next(reddit.subreddit("all").hot(limit=1))
            except Exception as e:
                # Mark entire job as failed if authentication fails
                jobs[job_id]["status"] = "error"
                for entry in jobs[job_id]["results"]:
                    entry["status"] = "error"
                    entry["error"] = f"Reddit authentication failed: {e}"
                return
            # Start processing
            jobs[job_id]["status"] = "processing"
            for i, url in enumerate(url_list):
                try:
                    (
                        tt,
                        tb,
                        tu,
                        comments,
                        media,
                        post_date,
                        sid,
                    ) = get_reddit_thread_data(reddit, url)
                    filename = generate_docx(
                        url,
                        tt,
                        tb,
                        tu,
                        comments,
                        media,
                        post_date,
                        job_dir,
                        sid,
                    )
                    jobs[job_id]["results"][i].update(
                        {
                            "status": "success",
                            "filename": filename,
                            "error": None,
                        }
                    )
                except Exception as e:
                    jobs[job_id]["results"][i].update(
                        {
                            "status": "error",
                            "filename": None,
                            "error": str(e),
                        }
                    )
                # Update completion count
                jobs[job_id]["completed"] = i + 1
                # Delay between requests except after last one
                if i < total - 1:
                    time.sleep(DELAY_SECONDS)
            # Create ZIP file
            zip_filename = f"{job_id}.zip"
            zip_path = os.path.join(job_dir, zip_filename)
            zip_folder(job_dir, zip_path)
            jobs[job_id]["zip_url"] = f"/api/download/{job_id}"
            jobs[job_id]["status"] = "completed"
        except Exception as e:
            jobs[job_id]["status"] = "error"
            # Propagate error to all result entries that haven't been marked
            for entry in jobs[job_id]["results"]:
                if entry["status"] == "pending":
                    entry["status"] = "error"
                    entry["error"] = str(e)
            jobs[job_id]["zip_url"] = None

    # Start background thread
    thread = threading.Thread(target=process_job, daemon=True)
    thread.start()

    return {"job_id": job_id}


@app.get("/api/download/{job_id}")
async def download_zip(job_id: str):
    """Serve the ZIP archive for a given job ID."""
    zip_file = os.path.join(BASE_JOBS_DIR, job_id, f"{job_id}.zip")
    if not os.path.exists(zip_file):
        raise HTTPException(status_code=404, detail="ZIP file not found.")
    return FileResponse(
        zip_file,
        media_type="application/zip",
        filename=f"{job_id}.zip",
    )


# -----------------------------------------------------------------------------
# Job status endpoint
#
# Allows clients to poll for the progress of a given job. Returns the job's
# current status, total number of URLs, how many have completed, per‑URL
# results, and the relative download path once available. Returns 404 if
# the job ID is unknown.

@app.get("/api/status/{job_id}")
async def get_status(job_id: str):
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job